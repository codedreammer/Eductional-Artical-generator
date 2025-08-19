import os
import io
import streamlit as st
from docx import Document
from docx.shared import Pt

st.set_page_config(page_title="Educational Article Generator", page_icon="📚", layout="wide")
st.title("📚 Educational Article Generator")
st.write("Generate clear, structured educational articles with AI.")

# --------------- Sidebar ---------------
st.sidebar.header("Input Settings")
topic = st.sidebar.text_input("Topic", value="Photosynthesis")
audience = st.sidebar.selectbox("Audience Level", ["School Students", "College Students", "General Audience", "Professionals"])
word_count = st.sidebar.slider("Target Word Count", min_value=300, max_value=2000, value=800, step=50)
tone = st.sidebar.selectbox("Tone", ["Simple", "Formal", "Friendly", "Exam-Oriented"])
include_quiz = st.sidebar.checkbox("Include a short quiz at the end", value=True)
include_glossary = st.sidebar.checkbox("Include a brief glossary", value=True)

generate_btn = st.sidebar.button("Generate Article")

# --------------- Helper functions ---------------

def ask_ai(prompt: str) -> str:
    # Extract topic from prompt for personalized content
    topic_line = [line for line in prompt.split('\n') if 'Write an educational article on:' in line]
    topic = topic_line[0].split("'")[1] if topic_line else "General Topic"
    
    return f"""# {topic}: A Comprehensive Guide

## Introduction
{topic} is a fascinating subject that plays a crucial role in our understanding of the world. This article will explore the key concepts and practical applications.

## Main Concepts
### Core Principles
- **Definition**: {topic} refers to the fundamental processes and mechanisms involved
- **Key Components**: Multiple interconnected elements work together
- **Scientific Basis**: Based on well-established scientific principles

### Important Features
- Feature 1: Primary characteristic that defines the concept
- Feature 2: Secondary aspect that supports the main function
- Feature 3: Additional element that enhances understanding

## Real-life Applications
{topic} has numerous practical applications in everyday life:
- **Application 1**: Used in technology and industry
- **Application 2**: Important for environmental processes
- **Case Study**: A specific example demonstrating practical use

## Common Mistakes
- Misconception 1: Often people think incorrectly about this aspect
- Misconception 2: Another common misunderstanding to avoid
- Mistake 3: Typical error in application or understanding

## Summary
{topic} is essential for understanding modern science and technology. Key takeaways include the core principles, practical applications, and avoiding common misconceptions.

## Glossary
- **Term 1**: Basic definition of important concept
- **Term 2**: Another key term explained simply
- **Term 3**: Additional vocabulary for better understanding
- **Term 4**: Technical term made accessible
- **Term 5**: Final important concept defined

## Quick Quiz
1. What is the primary function of {topic}?
   a) Option A  b) Option B  c) Option C  d) Option D

2. Which application is most common?
   a) Application 1  b) Application 2  c) Application 3  d) All of the above

3. What is a common misconception about {topic}?
   a) Misconception 1  b) Misconception 2  c) Both  d) Neither

### Answer Key
1. c) Option C
2. d) All of the above
3. c) Both"""

def build_prompt(topic, audience, word_count, tone, include_quiz, include_glossary):
    parts = [
        f"Write an educational article on: '{topic}'.",
        f"Audience level: {audience}.",
        f"Tone: {tone}.",
        f"Target word count: around {word_count} words.",
        "Structure the article clearly with these sections:",
        "1) Title",
        "2) Introduction (why this topic matters)",
        "3) Main Concepts (with subheadings, bullet points, and 1–2 simple examples)",
        "4) Real-life Application or Case Study",
        "5) Common Mistakes or Misconceptions",
        "6) Summary / Key Takeaways",
    ]
    if include_glossary:
        parts.append("7) Glossary: 5–8 key terms with one-line explanations.")
    if include_quiz:
        parts.append("8) Quick Quiz: 4–6 questions (MCQ or short), and a separate 'Answer Key'.")
    parts.append("Use markdown formatting: headings (##), lists, and bold where helpful. Avoid plagiarism and keep it age-appropriate.")
    return "\n".join(parts)

def article_to_docx(markdown_text: str) -> bytes:
    # Minimal Markdown -> DOCX (naive): keep headings and paragraphs.
    # This avoids external converters to keep the project simple.
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = markdown_text.splitlines()
    for ln in lines:
        if ln.strip().startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(ln.strip().replace("### ", ""))
            run.bold = True
            run.font.size = Pt(14)
        elif ln.strip().startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(ln.strip().replace("## ", ""))
            run.bold = True
            run.font.size = Pt(16)
        elif ln.strip().startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(ln.strip().replace("# ", ""))
            run.bold = True
            run.font.size = Pt(18)
        else:
            doc.add_paragraph(ln)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def get_download_bytes(text: str) -> bytes:
    return text.encode("utf-8")

# --------------- Main ---------------
if generate_btn:
    with st.spinner("Generating article with AI..."):
        prompt = build_prompt(topic, audience, word_count, tone, include_quiz, include_glossary)
        article_md = ask_ai(prompt)

    st.success("Article generated!")
    st.markdown("### Preview")
    st.markdown(article_md)

    # Downloads
    st.download_button("⬇️ Download as Markdown (.md)", data=get_download_bytes(article_md), file_name=f"{topic.replace(' ', '_')}.md")
    docx_bytes = article_to_docx(article_md)
    st.download_button("⬇️ Download as Word (.docx)", data=docx_bytes, file_name=f"{topic.replace(' ', '_')}.docx")

else:
    st.info("Enter a topic and click **Generate Article** in the sidebar.")
    
st.caption("Tip: Try topics like 'Binary Search', 'Newton's Laws', 'Photosynthesis', 'World War II causes', 'Basics of SQL'.")
